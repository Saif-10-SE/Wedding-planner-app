from datetime import datetime
from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(__file__).resolve().parent
OUT = ROOT / "Wedify_Project_Proposal.docx"


def add_title(doc, title, subtitle=None):
    p = doc.add_paragraph(title)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.runs[0]
    run.bold = True
    run.font.size = Pt(20)

    if subtitle:
        sp = doc.add_paragraph(subtitle)
        sp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        srun = sp.runs[0]
        srun.font.size = Pt(11)


def add_heading(doc, text, level=1):
    doc.add_heading(text, level=level)


def add_para(doc, text):
    doc.add_paragraph(text)


def add_bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style='List Bullet')


def add_numbered(doc, items):
    for item in items:
        doc.add_paragraph(item, style='List Number')


def build_document():
    doc = Document()

    add_title(
        doc,
        "Project Proposal",
        "Project Name: Wedify\nTitle: Luxury Desi Wedding Budget Calculator & Planning Platform\nProgram: BS Software Engineering / Computer Science\nPrepared For: University Academic Submission\nDate: " + datetime.now().strftime("%d %B %Y"),
    )

    doc.add_paragraph("\n")

    add_heading(doc, "Abstract", 1)
    add_para(doc,
        "Wedify is proposed as a premium digital planning platform designed to address the financial and managerial complexity of high-budget South Asian weddings. "
        "The project combines intelligent budgeting, structured planning workflows, vendor cost modeling, and analytics-oriented dashboards in a single integrated ecosystem. "
        "The proposed system emphasizes cultural alignment with Desi wedding ceremonies such as Mehndi, Barat, and Walima while introducing modern software engineering practices, data-driven decision support, and scalable system architecture. "
        "The academic and practical contribution of the project lies in formalizing wedding financial planning as a measurable and optimizable process rather than an ad hoc, manually coordinated activity."
    )

    add_heading(doc, "1. Introduction", 1)
    add_heading(doc, "1.1 Problem Statement", 2)
    add_para(doc,
        "Wedding planning in South Asian contexts is traditionally managed through fragmented communication channels, manual records, and socially driven decision-making. "
        "These methods create substantial uncertainty in cost forecasting, event-wise budget distribution, and vendor coordination. "
        "Families frequently face budget overruns due to hidden expenditures, incomplete early-stage planning, and lack of real-time visibility into cumulative financial commitments. "
        "The absence of a culturally contextualized, intelligent digital platform for elite wedding planning constitutes a clear software and process gap."
    )

    add_heading(doc, "1.2 Background: Wedding Planning Challenges in South Asian Culture", 2)
    add_para(doc,
        "South Asian weddings are multi-event, multi-stakeholder, and symbolically significant social systems. "
        "Unlike single-ceremony models, Desi weddings involve several major events, extensive guest management, custom décor requirements, and personalized hospitality standards. "
        "Cultural expectations regarding prestige, aesthetics, and family reputation intensify planning pressure and decision complexity. "
        "Consequently, planning quality and financial discipline become critical determinants of event success."
    )

    add_heading(doc, "1.3 Financial Planning Difficulties in High-Budget Weddings", 2)
    add_para(doc,
        "High-budget weddings involve large and highly variable cost components: venue tiers, per-head catering rates, décor themes, entertainment packages, logistics, media production, and contingency reserves. "
        "In traditional settings, these components are estimated independently and updated inconsistently, producing weak budget governance. "
        "Without dynamic calculation models, scenario-based comparison, and expenditure tracking, families face high risks of financial inefficiency and unplanned debt pressure."
    )

    add_heading(doc, "1.4 Need for an Intelligent Wedding Budgeting Platform", 2)
    add_para(doc,
        "The proposed platform is needed to operationalize financial planning through structured data input, rule-based computation, and transparent analytics. "
        "Wedify aims to provide event-wise budget simulation, real-time recalculation, vendor estimation intelligence, and dashboard-level visibility for decision-making. "
        "By embedding cultural event logic within a modern software architecture, the platform can reduce uncertainty, improve accountability, and raise planning quality for upper-class Desi households."
    )

    add_heading(doc, "2. Project Objectives", 1)
    add_heading(doc, "2.1 Primary Objectives", 2)
    add_bullets(doc, [
        "To design and implement an intelligent budgeting engine for multi-event South Asian weddings.",
        "To provide real-time cost recalculation based on user choices and dynamic constraints.",
        "To enable high-precision financial planning through dashboard visualizations and scenario analysis.",
        "To centralize planning workflows for venue, vendor, and expense management within one platform.",
    ])

    add_heading(doc, "2.2 Secondary Objectives", 2)
    add_bullets(doc, [
        "To improve user confidence through transparent cost breakdowns and comparative insights.",
        "To reduce planning overhead by integrating checklist and event lifecycle management.",
        "To support informed decision-making with visual analytics and category-level expenditure summaries.",
        "To establish a maintainable codebase suitable for academic demonstration and real-world extension.",
    ])

    add_heading(doc, "2.3 Long-Term Scalability Goals", 2)
    add_bullets(doc, [
        "To evolve into a cloud-enabled, account-based planning service for regional and global markets.",
        "To integrate AI-driven recommendations for vendors, décor concepts, and budget optimization.",
        "To support multilingual interfaces and broader cultural ceremony models beyond South Asia.",
        "To establish analytics-based forecasting for cost inflation and seasonal pricing variability.",
    ])

    add_heading(doc, "3. Scope of the Project", 1)
    add_heading(doc, "3.1 Included Scope", 2)
    add_bullets(doc, [
        "Multi-event planning for Mehndi, Barat, and Walima.",
        "Dynamic budget calculation by venue, guest count, and service category.",
        "Vendor price estimation and comparative cost visualization.",
        "Expense tracking modules and planning status workflows.",
        "Responsive luxury UI for desktop and mobile experiences.",
    ])

    add_heading(doc, "3.2 Excluded Scope", 2)
    add_bullets(doc, [
        "Direct payment gateway integration in initial release.",
        "Legally binding vendor contract automation.",
        "Physical logistics execution and on-ground event operations.",
        "Real-time IoT integrations for venue infrastructure.",
    ])

    add_heading(doc, "3.3 Target Audience", 2)
    add_para(doc,
        "The primary target audience is upper-class and high-net-worth Desi families who prioritize premium planning quality, financial transparency, and ceremonial excellence. "
        "Secondary users include wedding planners, event consultants, and venue coordinators seeking structured budgeting workflows."
    )

    add_heading(doc, "3.4 Functional Boundaries", 2)
    add_para(doc,
        "Wedify will act as a planning and decision-support platform rather than a transactional marketplace in the initial phase. "
        "System responsibility is limited to budgeting intelligence, planning workflows, visualization, and administration modules with controlled extensibility for future integrations."
    )

    add_heading(doc, "4. Literature Review", 1)
    add_heading(doc, "4.1 Existing Wedding Applications", 2)
    add_para(doc,
        "Current wedding applications largely emphasize invitation management, checklist reminders, and template-based planning. "
        "Most available products are optimized for Western single-day ceremonies and do not adequately model the multi-event structure of South Asian weddings."
    )

    add_heading(doc, "4.2 Budget Calculators in the Market", 2)
    add_para(doc,
        "Generic budget calculators provide static category forms and percentage-based totals, but they lack event-aware decomposition, tiered vendor intelligence, and culturally specific cost models. "
        "They rarely support iterative scenario planning for high-variance expenses such as décor and hospitality."
    )

    add_heading(doc, "4.3 Limitations of Current Platforms", 2)
    add_bullets(doc, [
        "Insufficient localization for South Asian ceremonies and spending patterns.",
        "Weak real-time recalculation and low transparency in assumptions.",
        "No integrated elite-focused UX strategy for premium market segments.",
        "Limited analytics depth and poor support for strategic cost optimization.",
    ])

    add_heading(doc, "4.4 Gap Analysis", 2)
    add_para(doc,
        "The review identifies a clear gap: a culturally aligned, intelligent, and analytics-driven platform for luxury Desi wedding budgeting and planning. "
        "Wedify addresses this gap by combining event-specific computation logic, premium UX design, structured planning modules, and extensible architecture for future AI integration."
    )

    add_heading(doc, "5. System Architecture", 1)
    add_heading(doc, "5.1 Overall Architecture Description", 2)
    add_para(doc,
        "The proposed architecture follows a layered service-oriented model: presentation layer, application service layer, domain/business logic layer, persistence layer, and deployment layer. "
        "The frontend consumes RESTful services from a Node.js/Express backend, while data is stored in a structured database with normalized entities for events, budgets, vendors, users, and expenses."
    )

    add_heading(doc, "5.2 Frontend Architecture", 2)
    add_para(doc,
        "The frontend is built with Next.js and React component architecture, organized into route-level pages, reusable UI components, context/state management, and visualization modules. "
        "Client-side interactivity includes dynamic form controls, live budget updates, and dashboard rendering with chart-based insights."
    )

    add_heading(doc, "5.3 Backend Architecture", 2)
    add_para(doc,
        "The backend follows a modular Express architecture with clear separation of controllers, services, repositories, and validation middleware. "
        "Business logic encapsulates event-wise budget rules, vendor pricing normalization, and recommendation interfaces. "
        "Authentication and role-based access control are managed at API gateway and middleware layers."
    )

    add_heading(doc, "5.4 Database Schema Overview", 2)
    add_para(doc,
        "The schema models users, wedding events, budget plans, expense items, vendor catalogs, quotations, and administrative metadata. "
        "Relational integrity (or document-level references in NoSQL variants) is used to maintain data consistency between event records and financial transactions."
    )

    add_heading(doc, "5.5 Hosting and Deployment Architecture", 2)
    add_para(doc,
        "The frontend is proposed for deployment on Vercel/Netlify with CDN-backed static assets and edge delivery. "
        "Backend services are deployable to container-friendly Node hosting environments with HTTPS termination, environment-based configuration, and CI/CD pipelines from GitHub. "
        "Database hosting can be managed via cloud providers with backup, monitoring, and encryption-at-rest policies."
    )

    add_heading(doc, "5.6 Diagram Descriptions (Text Form)", 2)
    add_heading(doc, "System Architecture Diagram Description", 3)
    add_para(doc,
        "The diagram contains five horizontal layers. At the top, users interact with the Web UI. Requests move to the Next.js frontend layer, then to the REST API layer (Node.js/Express). "
        "The API communicates with the business rules module for budget computation and vendor logic. Finally, persistence operations are performed in the database layer. "
        "Cross-cutting services such as authentication, logging, and monitoring connect vertically across all layers."
    )

    add_heading(doc, "Data Flow Diagram (DFD Level 0) Description", 3)
    add_para(doc,
        "At Level 0, the Wedify system is represented as a single process receiving inputs from Users and Admins. "
        "Users provide event and budget parameters; the system returns budget outputs, reports, and planning updates. "
        "Admins provide vendor and pricing data; the system stores updates and returns management summaries."
    )

    add_heading(doc, "Data Flow Diagram (DFD Level 1) Description", 3)
    add_para(doc,
        "Level 1 decomposes the system into processes: User Management, Event Planning, Budget Engine, Vendor Estimation, Expense Tracking, and Dashboard Analytics. "
        "Each process exchanges data with specific data stores such as User Store, Event Store, Vendor Store, and Expense Store. "
        "The Budget Engine consumes event specifications and vendor prices, then writes computed plans to Budget Store and sends computed metrics to Dashboard Analytics."
    )

    add_heading(doc, "ER Diagram Description", 3)
    add_para(doc,
        "The ER model includes entities: User, WeddingPlan, Event, BudgetItem, Expense, Vendor, Quote, and AdminAction. "
        "User has one-to-many relationship with WeddingPlan; WeddingPlan has one-to-many with Event and BudgetItem; Event links to multiple Vendor quotes; Expenses map to BudgetItems for variance analysis. "
        "Primary keys are surrogate IDs; foreign keys enforce referential integrity across planning and financial records."
    )

    add_heading(doc, "Use Case Diagram Description", 3)
    add_para(doc,
        "Actors: Guest User, Registered User, Planner User, and Admin. Core use cases include authentication, event setup, budget simulation, vendor estimation, expense logging, report generation, and admin catalog management. "
        "Include/extend relationships model optional flows such as AI décor suggestion and comparative scenario analysis."
    )

    add_heading(doc, "Sequence Diagram Description", 3)
    add_para(doc,
        "A typical sequence starts when User submits event and budget preferences through UI. Frontend sends API request to Budget Service. Service retrieves pricing references from Vendor Repository and policy rules from Budget Rule Engine. "
        "Computed totals are persisted and returned to frontend. Frontend updates charts and dashboard widgets in real time and logs user interaction analytics."
    )

    add_heading(doc, "6. Tools and Technologies Used (Detailed)", 1)
    add_heading(doc, "6.1 Frontend", 2)
    add_para(doc,
        "React / Next.js: Selected for component-driven architecture, routing flexibility, rendering optimization, and strong ecosystem maturity. "
        "Next.js additionally supports server-side rendering and deployment efficiency, suitable for high-quality user experiences."
    )
    add_para(doc,
        "Tailwind CSS: Chosen for systematic utility-first styling, rapid UI iteration, design-token consistency, and maintainable luxury interface implementation."
    )
    add_para(doc,
        "Framer Motion: Selected for premium animation choreography and interaction transitions that enhance perceived quality and user engagement without compromising maintainability."
    )
    add_para(doc,
        "Chart Libraries (e.g., Recharts/Chart.js): Chosen to represent budget proportions, trend lines, and comparative cost analytics in interpretable visual formats."
    )

    add_heading(doc, "6.2 Backend", 2)
    add_para(doc,
        "Node.js / Express: Chosen for asynchronous I/O performance, broad package support, and efficient implementation of RESTful API layers. "
        "Express provides lightweight middleware composition suitable for modular service-oriented backend development."
    )

    add_heading(doc, "6.3 Database", 2)
    add_para(doc,
        "MongoDB / PostgreSQL: The design supports either document-oriented flexibility (MongoDB) or relational consistency and complex querying (PostgreSQL). "
        "Selection can be finalized according to reporting complexity and transactional strictness requirements."
    )

    add_heading(doc, "6.4 Version Control", 2)
    add_para(doc,
        "Git / GitHub: Chosen for distributed versioning, branch-based collaboration, traceable commits, and CI/CD integration readiness."
    )

    add_heading(doc, "6.5 Deployment", 2)
    add_para(doc,
        "Vercel / Netlify: Selected for straightforward frontend deployment, CDN-backed distribution, and environment-variable based workflow for staged releases."
    )

    add_heading(doc, "6.6 Testing", 2)
    add_para(doc,
        "Manual Testing: required for validating UX quality, content correctness, and ceremony-specific workflows. "
        "Unit Testing: required for validating isolated business logic, especially budget calculations. "
        "UI Testing: required for interaction reliability, responsive behavior, and regression safety."
    )

    add_heading(doc, "7. Functional Requirements", 1)
    add_numbered(doc, [
        "User registration and login with secure credential management.",
        "Event selection and configuration for Mehndi, Barat, and Walima.",
        "Dynamic budget calculation with category-level and event-level totals.",
        "Vendor price estimation with selectable package tiers.",
        "Real-time cost adjustment upon parameter updates.",
        "Expense tracking against planned budgets with variance visibility.",
        "Visualization dashboard for totals, allocations, and comparative insights.",
        "Administrative panel for vendor catalog, pricing rules, and content control.",
        "AI-assisted décor idea generation aligned with event type and budget range.",
    ])

    add_heading(doc, "8. Non-Functional Requirements", 1)
    add_bullets(doc, [
        "Performance: sub-second response for standard budget recalculations and optimized rendering under common load.",
        "Security: encrypted transport (HTTPS), secure authentication, input validation, and role-based authorization.",
        "Scalability: horizontally extensible service architecture with cache-ready API design.",
        "Responsiveness: consistent functionality across mobile, tablet, and desktop viewports.",
        "Usability: low-friction navigation, transparent data presentation, and guided multi-step workflows.",
        "Availability: production-grade uptime targets through cloud deployment and monitoring.",
        "Maintainability: modular code organization, documentation, test coverage, and standardized conventions.",
    ])

    add_heading(doc, "9. Database Design", 1)
    add_heading(doc, "9.1 Entities", 2)
    add_bullets(doc, [
        "User",
        "WeddingPlan",
        "Event",
        "BudgetItem",
        "Expense",
        "Vendor",
        "VendorPackage",
        "Quote",
        "AdminUser",
        "AuditLog",
    ])

    add_heading(doc, "9.2 Key Attributes", 2)
    add_bullets(doc, [
        "User: user_id, full_name, email, phone, password_hash, role, created_at",
        "WeddingPlan: plan_id, user_id, title, target_budget, status, created_at",
        "Event: event_id, plan_id, event_type, event_date, guest_count, venue_id",
        "BudgetItem: item_id, event_id, category, planned_cost, actual_cost",
        "Expense: expense_id, item_id, amount, payment_status, recorded_at",
        "Vendor: vendor_id, vendor_type, name, location, rating",
        "Quote: quote_id, vendor_id, event_id, package_id, quoted_amount",
    ])

    add_heading(doc, "9.3 Relationships", 2)
    add_bullets(doc, [
        "User 1..* WeddingPlan",
        "WeddingPlan 1..* Event",
        "Event 1..* BudgetItem",
        "BudgetItem 1..* Expense",
        "Vendor 1..* Quote",
        "Event 1..* Quote",
    ])

    add_heading(doc, "9.4 Keys", 2)
    add_para(doc,
        "Primary keys are synthetic identifiers for each entity. Foreign keys enforce references: WeddingPlan.user_id -> User.user_id, Event.plan_id -> WeddingPlan.plan_id, BudgetItem.event_id -> Event.event_id, Expense.item_id -> BudgetItem.item_id, Quote.vendor_id -> Vendor.vendor_id, Quote.event_id -> Event.event_id."
    )

    add_heading(doc, "9.5 Normalization", 2)
    add_para(doc,
        "The relational model is designed to satisfy Third Normal Form (3NF) by eliminating partial and transitive dependencies. "
        "Category and package data are isolated from transactional expense records to reduce redundancy and improve update consistency."
    )

    add_heading(doc, "10. User Interface Design Philosophy", 1)
    add_heading(doc, "10.1 Luxury UI Direction", 2)
    add_para(doc,
        "The interface follows a premium visual language characterized by high contrast hierarchy, ceremonial aesthetics, and refined typography."
    )

    add_heading(doc, "10.2 Color Psychology", 2)
    add_para(doc,
        "A jewel-tone palette (maroon, emerald, midnight, warm gold accents) is selected to communicate prestige, trust, celebration, and financial seriousness."
    )

    add_heading(doc, "10.3 Animation Strategy", 2)
    add_para(doc,
        "Motion is used functionally for feedback, transitions, and progressive disclosure, avoiding decorative overload while preserving luxury perception."
    )

    add_heading(doc, "10.4 Elite Targeting UX", 2)
    add_para(doc,
        "UX patterns prioritize confidence and control: transparent cost cards, immediate recalculation cues, premium dashboard composition, and reduced cognitive friction for high-stakes decisions."
    )

    add_heading(doc, "11. Implementation Plan", 1)
    add_heading(doc, "Phase 1: Research (Weeks 1-2)", 2)
    add_bullets(doc, [
        "Requirement elicitation and stakeholder interviews.",
        "Market analysis and feature benchmarking.",
        "Data model conceptualization and feasibility assessment.",
    ])

    add_heading(doc, "Phase 2: Design (Weeks 3-4)", 2)
    add_bullets(doc, [
        "Information architecture and user journey mapping.",
        "Wireframes, UI prototypes, and design system tokens.",
        "Architecture and API contract finalization.",
    ])

    add_heading(doc, "Phase 3: Development (Weeks 5-9)", 2)
    add_bullets(doc, [
        "Frontend implementation of core pages and reusable components.",
        "Backend API development and business-rule integration.",
        "Database schema implementation and service connectivity.",
    ])

    add_heading(doc, "Phase 4: Testing (Weeks 10-11)", 2)
    add_bullets(doc, [
        "Unit testing for budget logic and data integrity.",
        "Manual and UI testing for workflow completeness.",
        "Performance tuning and bug resolution.",
    ])

    add_heading(doc, "Phase 5: Deployment (Week 12)", 2)
    add_bullets(doc, [
        "Production deployment setup and environment hardening.",
        "Final acceptance testing and documentation handover.",
        "Post-deployment monitoring and iteration planning.",
    ])

    add_heading(doc, "12. Risk Analysis", 1)
    add_heading(doc, "12.1 Technical Risks", 2)
    add_bullets(doc, [
        "Complexity in maintaining consistent budget logic across event permutations.",
        "Integration overhead between frontend state and backend validation.",
        "Potential performance degradation with large scenario datasets.",
    ])

    add_heading(doc, "12.2 Financial Risks", 2)
    add_bullets(doc, [
        "Cost inflation in vendor market data reducing estimate accuracy.",
        "Unexpected hosting or scaling expenses in later adoption stages.",
    ])

    add_heading(doc, "12.3 Scaling Risks", 2)
    add_bullets(doc, [
        "Concurrency pressure under expanded user load.",
        "Data growth affecting query performance and reporting latency.",
    ])

    add_heading(doc, "12.4 Mitigation Strategies", 2)
    add_bullets(doc, [
        "Modular rule engine with test coverage for deterministic budget computation.",
        "Caching and pagination for analytics-heavy views.",
        "Cloud monitoring, alerting, and capacity planning mechanisms.",
        "Regular pricing model recalibration with admin-governed updates.",
    ])

    add_heading(doc, "13. Future Enhancements", 1)
    add_bullets(doc, [
        "AI-based vendor recommendations from event profile, budget, and quality signals.",
        "Predictive budgeting using historical trends and seasonal market behavior.",
        "Multi-event scenario comparison with optimization scoring.",
        "International wedding support through localization and currency frameworks.",
    ])

    add_heading(doc, "14. Conclusion", 1)
    add_para(doc,
        "Wedify represents an academically grounded and industry-relevant response to the financial and operational complexity of luxury South Asian weddings. "
        "By combining culturally contextual planning workflows with rigorous software architecture, the proposed platform transforms wedding budgeting from reactive estimation to proactive decision intelligence. "
        "Its socio-economic value lies in improving financial transparency, reducing wasteful expenditure, and increasing planning confidence for families undertaking high-investment ceremonies. "
        "From a software engineering perspective, the project demonstrates the integration of modern web technologies, scalable system design, and user-centered product strategy in a domain with substantial unmet digital needs."
    )

    add_heading(doc, "References (Indicative for Final Dissertation)", 1)
    add_numbered(doc, [
        "Pressman, R. S., & Maxim, B. R. Software Engineering: A Practitioner's Approach.",
        "Sommerville, I. Software Engineering.",
        "Nielsen, J. Usability Engineering.",
        "Official documentation for React, Next.js, Node.js, and selected database platform.",
    ])

    doc.save(OUT)
    print(f"Created: {OUT}")


if __name__ == "__main__":
    build_document()

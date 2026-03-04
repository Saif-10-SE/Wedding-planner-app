from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = Path(__file__).resolve().parent
OUT = ROOT / "Wedify_Clone_Rebuild_Prompt.docx"

PROMPT_TEXT = """Build a complete production-ready Next.js (Pages Router) web application named Wedify (luxury Desi wedding planning platform) that is a functional clone of an elite Lahore wedding planner system.

You must generate all code files, data files, config files, and styles from scratch in one workspace.

1) Project Goal
Create a premium, responsive web app for upper-class Desi families to:
- Discover and filter wedding venues
- Compare venues side-by-side
- Build multi-event wedding budgets
- Save favorites
- Track planning checklist progress
- Browse testimonials, vendors, and gallery inspiration

It must feel high-end, emotional, and ceremonial: “luxury couture wedding meets modern SaaS intelligence.”

2) Tech Stack (strict)
- Next.js 14, React 18, Pages Router
- Tailwind CSS 3
- Lucide React icons
- No TypeScript (JavaScript only)
- Use local static data files (no backend DB required)
- Use browser localStorage for persistence

3) Required Routes
Create these pages:
- /
- /marquees
- /marquees/[slug]
- /calculator
- /checklist
- /compare
- /favorites
- /vendors
- /gallery
- /testimonials
- Include _app.js and _document.js

4) Required Reusable Components
Build reusable components for:
- Navbar (desktop + mobile dropdown, compare/favorite counters, CTA)
- Footer
- MarqueeCard
- VendorCard
- TestimonialCard
- ImageGallery
- InquiryForm
- FavoriteButton
- CompareButton
- CountdownTimer
- SearchModal
- WeddingDateModal
- NotificationToast

5) Global State (must use React Context)
Create a global context provider with:
- favorites and toggleFavorite
- compareList and toggleCompare (max 4 items)
- global searchQuery, searchResults, isSearchOpen, performSearch
- weddingDate, setWeddingDate, getDaysUntilWedding
- recentlyViewed and addToRecentlyViewed
- inquiryCart, addToInquiry, removeFromInquiry
- toast notifications via showNotification

Persist these localStorage keys:
- wedding_favorites
- compare_list
- wedding_date
- recently_viewed
- inquiry_cart
- checklist keys: checklist_completed, checklist_notes

6) Data Modules (static mock data)
Create structured data files for:
- Venues/marquees (10+ venues)
- Vendors (multiple categories)
- Testimonials
- Checklist timeline
- Marquee categories + metadata (A+, A, B, C)
- Pricing structures:
  - per-head ranges
  - hall rental
  - menu packages
  - decor packages
  - amenities
  - contact details
Use realistic Lahore-style wedding data.

7) Budget Calculator Requirements (important)
Implement a strong multi-event calculator:
- Events: Mehndi, Barat, Walima (toggleable)
- Per-event guest ratios
- Venue-based pricing impact
- Menu category selection with presets
- Live counters/add-ons
- Decor, photography, entertainment, transport, invitations
- Additional custom budget input
- Dynamic recalculation in real time
- Breakdown view + total + per-guest cost
- Pie chart visualization
- Budget optimization suggestions
- Optional confetti on finalized interactions
- Maintain smooth UX and no calculation lag

8) Design System (must match luxury feel)
Use a jewel-tone premium palette:
- Primary deep maroon family
- Warm gold accent
- Emerald sections
- Midnight/indigo sections
- Wine/plum accents

UI requirements:
- Soft gradients, glow accents, glass cards where suitable
- Section identity colors (different page hero moods)
- Paisley/noise-like subtle texture overlays
- Premium serif + sans pairing
- High readability and clear hierarchy
- Avoid cheap/cluttered visuals

9) Page-Specific Visual Identity
- Home: cinematic luxury landing with section-based color identities
- Compare: midnight analytical theme
- Favorites: wine/rose emotional theme
- Gallery: maroon + gold editorial theme
- Testimonials: emerald trust/theme
- Vendors: midnight professional directory theme
- Checklist: emerald planning calm theme
- Calculator: royal maroon dashboard tone

10) UX/Behavior Requirements
- Fully responsive (mobile/tablet/desktop)
- Sticky filter bars where appropriate
- Smooth hover states and micro-interactions
- No broken routes
- No broken images (use valid Unsplash URLs)
- All cards/actions must be clickable and meaningful
- Empty states for favorites/compare/search results

11) Config & Setup Files
Generate:
- package.json
- next.config.js
- tailwind.config.js
- postcss.config.js
- jsconfig.json with alias mapping for @/
- styles/globals.css with custom utility classes/animations

12) Quality Bar
- Clean file/folder organization
- Reusable components, no unnecessary duplication
- No placeholder lorem text in core UX
- No TypeScript, no backend dependency
- Production build must pass (next build)
- Include concise README with run instructions

13) Deliverables
Output all source code files in full (or create them directly in workspace), including:
- all pages
- all components
- context provider
- all data modules
- full Tailwind and global CSS theme system
- README
- working app with no missing imports

14) Local Run Commands
Use:
- install: npm install
- dev: node node_modules/next/dist/bin/next dev -p 3000
- build: npx next build

Now generate the complete project with this exact structure and behavior. Do not provide a partial scaffold. Build the full, functional clone in one go."""


def build_doc():
    doc = Document()

    title = doc.add_paragraph("Wedify Clone Rebuild Prompt")
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].bold = True
    title.runs[0].font.size = Pt(18)

    subtitle = doc.add_paragraph("Copy this prompt and paste it into GitHub Copilot Chat to recreate the same project.")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("")
    body = doc.add_paragraph(PROMPT_TEXT)
    for run in body.runs:
        run.font.size = Pt(11)

    doc.save(OUT)
    print(f"Created: {OUT}")


if __name__ == "__main__":
    build_doc()

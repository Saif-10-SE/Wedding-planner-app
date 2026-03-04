from datetime import datetime
from pathlib import Path

import matplotlib.pyplot as plt
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


ROOT = Path(__file__).resolve().parent
OUT_DOC = ROOT / "WeddingPlanner_User_Manual.docx"
DIAGRAM_DIR = ROOT / "docs" / "diagrams"


def add_box(ax, x, y, w, h, text, face="#f4efe6", edge="#6b0f18", text_color="#1f2937", fontsize=9):
    box = FancyBboxPatch(
        (x, y), w, h,
        boxstyle="round,pad=0.02,rounding_size=0.02",
        linewidth=1.6,
        edgecolor=edge,
        facecolor=face,
    )
    ax.add_patch(box)
    ax.text(x + w / 2, y + h / 2, text, ha="center", va="center", fontsize=fontsize, color=text_color, wrap=True)


def add_arrow(ax, x1, y1, x2, y2, color="#374151"):
    arrow = FancyArrowPatch((x1, y1), (x2, y2), arrowstyle="->", mutation_scale=12, linewidth=1.4, color=color)
    ax.add_patch(arrow)


def diagram_architecture(path: Path):
    fig, ax = plt.subplots(figsize=(12, 7))
    ax.set_xlim(0, 1)
    ax.set_ylim(0, 1)
    ax.axis("off")

    ax.text(0.5, 0.95, "Lahore Shaadi - System Architecture", ha="center", va="center", fontsize=16, fontweight="bold")

    add_box(ax, 0.08, 0.77, 0.84, 0.11, "Presentation Layer\nNext.js Pages + Reusable Components + Tailwind UI", face="#fef3c7", edge="#b45309")
    add_box(ax, 0.08, 0.60, 0.84, 0.11, "State Layer\nWeddingContext (favorites, compare list, search, inquiry cart, wedding date, notifications)", face="#dbeafe", edge="#1d4ed8")
    add_box(ax, 0.08, 0.43, 0.84, 0.11, "Domain Data Layer\nData modules: marquees, vendors, testimonials, checklist, menu", face="#dcfce7", edge="#15803d")
    add_box(ax, 0.08, 0.26, 0.84, 0.11, "Persistence Layer\nBrowser localStorage (favorites, compare list, inquiry cart, date, checklist progress)", face="#fce7f3", edge="#9d174d")
    add_box(ax, 0.08, 0.09, 0.84, 0.11, "Runtime Layer\nNode.js + Next.js dev/build pipeline", face="#ede9fe", edge="#5b21b6")

    add_arrow(ax, 0.50, 0.77, 0.50, 0.71)
    add_arrow(ax, 0.50, 0.60, 0.50, 0.54)
    add_arrow(ax, 0.50, 0.43, 0.50, 0.37)
    add_arrow(ax, 0.50, 0.26, 0.50, 0.20)

    fig.tight_layout()
    fig.savefig(path, dpi=220)
    plt.close(fig)


def diagram_user_workflows(path: Path):
    fig, ax = plt.subplots(figsize=(13, 8))
    ax.set_xlim(0, 1)
    ax.set_ylim(0, 1)
    ax.axis("off")

    ax.text(0.5, 0.96, "Primary User Workflows", ha="center", va="center", fontsize=16, fontweight="bold")

    add_box(ax, 0.04, 0.78, 0.18, 0.12, "Home\nDiscover tools & venues", face="#fef9c3", edge="#ca8a04")
    add_box(ax, 0.28, 0.78, 0.18, 0.12, "Marquees Listing\nFilter & sort", face="#fee2e2", edge="#be123c")
    add_box(ax, 0.52, 0.78, 0.18, 0.12, "Venue Detail\nGallery, tabs, inquiry", face="#d1fae5", edge="#047857")
    add_box(ax, 0.76, 0.78, 0.18, 0.12, "Actions\nFavorite / Compare / Budget", face="#dbeafe", edge="#1e40af")

    add_box(ax, 0.10, 0.52, 0.22, 0.14, "Budget Calculator\nMulti-event estimate\nExport/print proposal", face="#fce7f3", edge="#9d174d")
    add_box(ax, 0.39, 0.52, 0.22, 0.14, "Checklist\nTrack tasks + date\nSave progress", face="#dcfce7", edge="#15803d")
    add_box(ax, 0.68, 0.52, 0.22, 0.14, "Compare/Favorites\nShortlist decisions", face="#ede9fe", edge="#5b21b6")

    add_box(ax, 0.18, 0.24, 0.20, 0.14, "Vendors\nSearch & contact", face="#ffedd5", edge="#c2410c")
    add_box(ax, 0.42, 0.24, 0.20, 0.14, "Gallery\nBrowse inspiration", face="#e0f2fe", edge="#0369a1")
    add_box(ax, 0.66, 0.24, 0.20, 0.14, "Testimonials\nValidate choices", face="#ecfccb", edge="#4d7c0f")

    add_arrow(ax, 0.22, 0.84, 0.28, 0.84)
    add_arrow(ax, 0.46, 0.84, 0.52, 0.84)
    add_arrow(ax, 0.70, 0.84, 0.76, 0.84)

    add_arrow(ax, 0.82, 0.78, 0.21, 0.66)
    add_arrow(ax, 0.82, 0.78, 0.50, 0.66)
    add_arrow(ax, 0.82, 0.78, 0.79, 0.66)

    add_arrow(ax, 0.21, 0.52, 0.28, 0.38)
    add_arrow(ax, 0.50, 0.52, 0.52, 0.38)
    add_arrow(ax, 0.79, 0.52, 0.76, 0.38)

    fig.tight_layout()
    fig.savefig(path, dpi=220)
    plt.close(fig)


def diagram_state_flow(path: Path):
    fig, ax = plt.subplots(figsize=(12, 7))
    ax.set_xlim(0, 1)
    ax.set_ylim(0, 1)
    ax.axis("off")

    ax.text(0.5, 0.95, "Context + Persistence Data Flow", ha="center", va="center", fontsize=16, fontweight="bold")

    add_box(ax, 0.05, 0.62, 0.27, 0.22, "UI Components\nFavoriteButton\nCompareButton\nSearchModal\nWeddingDateModal\nChecklist Page", face="#fef3c7", edge="#b45309")
    add_box(ax, 0.38, 0.62, 0.24, 0.22, "WeddingContext\nReact State + Actions\nshowNotification()", face="#dbeafe", edge="#1d4ed8")
    add_box(ax, 0.68, 0.62, 0.27, 0.22, "localStorage Keys\nwedding_favorites\ncompare_list\ninquiry_cart\nwedding_date\nchecklist_completed", face="#fce7f3", edge="#9d174d")

    add_box(ax, 0.22, 0.24, 0.24, 0.20, "Derived Selectors\ngetFavoriteVenues()\ngetRecentlyViewedVenues()\ngetDaysUntilWedding()", face="#dcfce7", edge="#15803d")
    add_box(ax, 0.54, 0.24, 0.24, 0.20, "Route Rendering\n/favorites\n/compare\n/checklist\n/marquees/[slug]", face="#ede9fe", edge="#5b21b6")

    add_arrow(ax, 0.32, 0.73, 0.38, 0.73)
    add_arrow(ax, 0.62, 0.73, 0.68, 0.73)
    add_arrow(ax, 0.68, 0.69, 0.62, 0.69)
    add_arrow(ax, 0.50, 0.62, 0.34, 0.44)
    add_arrow(ax, 0.50, 0.62, 0.66, 0.44)

    fig.tight_layout()
    fig.savefig(path, dpi=220)
    plt.close(fig)


def diagram_budget_pipeline(path: Path):
    fig, ax = plt.subplots(figsize=(13, 7))
    ax.set_xlim(0, 1)
    ax.set_ylim(0, 1)
    ax.axis("off")

    ax.text(0.5, 0.94, "Budget Calculator Computation Pipeline", ha="center", va="center", fontsize=16, fontweight="bold")

    add_box(ax, 0.03, 0.60, 0.18, 0.20, "Inputs\nVenue\nEvents\nGuest counts\nMenu selections", face="#ffedd5", edge="#c2410c")
    add_box(ax, 0.25, 0.60, 0.18, 0.20, "Menu Engine\ncalculateEventMenuPerHead\ncalculateCountersCost\ncalculateStartersCost", face="#fde68a", edge="#a16207")
    add_box(ax, 0.47, 0.60, 0.18, 0.20, "Service Costs\nDecor\nPhotography\nEntertainment\nTransport\nInvitations", face="#d1fae5", edge="#047857")
    add_box(ax, 0.69, 0.60, 0.18, 0.20, "Venue Charges\nPer-head base\nHall rental\nCategory constraints", face="#dbeafe", edge="#1e40af")

    add_box(ax, 0.25, 0.26, 0.22, 0.20, "Aggregation\nPer-event subtotal\nCombined subtotal\nTax + contingency", face="#fce7f3", edge="#9d174d")
    add_box(ax, 0.55, 0.26, 0.22, 0.20, "Outputs\nGrand Total\nPer Guest Cost\nPie chart share\nOptimization hints", face="#ede9fe", edge="#5b21b6")

    add_arrow(ax, 0.21, 0.70, 0.25, 0.70)
    add_arrow(ax, 0.43, 0.70, 0.47, 0.70)
    add_arrow(ax, 0.65, 0.70, 0.69, 0.70)
    add_arrow(ax, 0.34, 0.60, 0.35, 0.46)
    add_arrow(ax, 0.56, 0.60, 0.36, 0.46)
    add_arrow(ax, 0.78, 0.60, 0.38, 0.46)
    add_arrow(ax, 0.47, 0.36, 0.55, 0.36)

    fig.tight_layout()
    fig.savefig(path, dpi=220)
    plt.close(fig)


def write_title(doc: Document):
    title = doc.add_paragraph("Lahore Shaadi - Comprehensive User Manual & Technical Workflow Report")
    title.runs[0].font.size = Pt(20)
    title.runs[0].font.bold = True
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_paragraph(f"Generated on: {datetime.now().strftime('%d %B %Y, %H:%M')}\nVersion: 1.0")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_heading(doc: Document, text: str, level: int = 1):
    doc.add_heading(text, level=level)


def add_paragraph(doc: Document, text: str):
    doc.add_paragraph(text)


def add_bullets(doc: Document, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def add_numbered(doc: Document, items):
    for item in items:
        doc.add_paragraph(item, style="List Number")


def add_tech_stack_table(doc: Document):
    table = doc.add_table(rows=1, cols=4)
    hdr = table.rows[0].cells
    hdr[0].text = "Layer"
    hdr[1].text = "Technology"
    hdr[2].text = "Version"
    hdr[3].text = "Purpose"

    rows = [
        ("Frontend Framework", "Next.js (Pages Router)", "14.x", "Routing, SSR/static rendering, build pipeline"),
        ("UI Library", "React", "18.2", "Component architecture and stateful UI"),
        ("Styling", "Tailwind CSS", "3.3", "Utility-first design system with custom tokens"),
        ("Icons", "Lucide React", "0.294", "Consistent iconography across pages"),
        ("Runtime", "Node.js", "18+ recommended", "Development server and production build execution"),
        ("Build Tools", "PostCSS + Autoprefixer", "8.4 / 10.4", "CSS processing and compatibility"),
        ("Persistence", "Browser localStorage", "N/A", "Client-side persistence for user planning state"),
    ]

    for r in rows:
        row = table.add_row().cells
        for i, col in enumerate(r):
            row[i].text = col


def add_routes_table(doc: Document):
    table = doc.add_table(rows=1, cols=3)
    hdr = table.rows[0].cells
    hdr[0].text = "Route"
    hdr[1].text = "File"
    hdr[2].text = "Primary Workflow"

    rows = [
        ("/", "pages/index.js", "Homepage discovery, hero CTA, featured sections, planning entry points"),
        ("/marquees", "pages/marquees/index.js", "Venue exploration with filters, sorting, and view mode"),
        ("/marquees/[slug]", "pages/marquees/[slug].js", "Venue deep-dive: gallery, packages, amenities, actions"),
        ("/calculator", "pages/calculator.js", "Multi-event budgeting, per-head calculations, optimization hints"),
        ("/checklist", "pages/checklist.js", "Task tracking, date countdown, export progress"),
        ("/compare", "pages/compare.js", "Side-by-side comparison of selected venues"),
        ("/favorites", "pages/favorites.js", "Saved venue shortlist and quick compare path"),
        ("/vendors", "pages/vendors.js", "Vendor search, category filtering, sorting, contact actions"),
        ("/gallery", "pages/gallery.js", "Wedding photo browsing by venue and event type"),
        ("/testimonials", "pages/testimonials.js", "Review browsing, rating/guest sorting, trust signals"),
    ]

    for r in rows:
        row = table.add_row().cells
        for i, col in enumerate(r):
            row[i].text = col


def add_workflows(doc: Document):
    add_heading(doc, "End-to-End Workflows", level=1)

    add_heading(doc, "Workflow 1: Discover and Shortlist Venues", level=2)
    add_numbered(doc, [
        "User lands on homepage and navigates to All Marquees.",
        "User applies filters (category, area, min capacity, max budget) and sort preference.",
        "User opens a venue detail page for deep inspection.",
        "User adds venue to Favorites and/or Compare list.",
        "User repeats for multiple venues, then goes to Compare page for decision support.",
    ])

    add_heading(doc, "Workflow 2: Compare Finalists", level=2)
    add_numbered(doc, [
        "Compare page loads selected venue slugs and resolves full records from data.",
        "User changes any selector to swap a candidate venue.",
        "System renders row-wise comparison: rating, capacity, pricing, packages, amenities, contact.",
        "User proceeds either to venue detail page or budget calculator for final viability checks.",
    ])

    add_heading(doc, "Workflow 3: Build Wedding Budget", level=2)
    add_numbered(doc, [
        "User selects venue and event set (e.g., mehndi/barat/valima).",
        "System initializes event presets for menu, counters, and guest ratio defaults.",
        "User customizes menus, starter quantities, counters, and add-on services.",
        "Calculator aggregates per-event and shared costs, then applies tax/contingency logic.",
        "UI updates animated totals, pie chart split, and optional optimization suggestions.",
    ])

    add_heading(doc, "Workflow 4: Manage Planning Checklist", level=2)
    add_numbered(doc, [
        "User sets wedding date via modal; countdown appears automatically.",
        "User checks/unchecks tasks across period buckets (12 months to day-before).",
        "System persists task completion and notes in localStorage.",
        "User exports checklist as JSON for offline backup or sharing.",
    ])

    add_heading(doc, "Workflow 5: Vendor Selection", level=2)
    add_numbered(doc, [
        "User searches vendors by text, category, and sorting criteria.",
        "User views featured vendors when no restrictive filter is applied.",
        "User switches between grid/list views based on review needs.",
        "User initiates direct contact (tel/mail) from vendor cards or list rows.",
    ])

    add_heading(doc, "Workflow 6: Gallery and Testimonials for Validation", level=2)
    add_numbered(doc, [
        "User opens Gallery for inspiration, filtered by venue or event type.",
        "User reviews testimonials and sorts by recency/rating/guest count.",
        "User uses social proof to validate final venue/vendor choices.",
    ])


def add_context_model(doc: Document):
    add_heading(doc, "WeddingContext State Model", level=1)
    add_bullets(doc, [
        "favorites: saved venue slugs with toggleFavorite() and isFavorite().",
        "compareList: up to 4 venue slugs with toggleCompare() limit enforcement.",
        "searchQuery/searchResults/isSearchOpen: global search modal control.",
        "weddingDate: selected date with getDaysUntilWedding() utility.",
        "recentlyViewed: capped recency list from venue detail visits.",
        "inquiryCart: list of inquiry intents keyed by venue slug.",
        "notification: ephemeral toast state managed by showNotification().",
    ])

    add_heading(doc, "localStorage Keys and Purpose", level=2)
    add_bullets(doc, [
        "wedding_favorites: persists shortlist across sessions.",
        "compare_list: preserves comparison set.",
        "inquiry_cart: stores venues queued for inquiry.",
        "wedding_date: stores selected wedding date in ISO format.",
        "recently_viewed: retains recent venue history.",
        "checklist_completed/checklist_notes: stores planning task progress and notes.",
    ])


def add_devops(doc: Document):
    add_heading(doc, "Setup and Operations", level=1)
    add_heading(doc, "Prerequisites", level=2)
    add_bullets(doc, [
        "Windows/macOS/Linux with Node.js 18+.",
        "npm package manager.",
        "Modern browser for local testing.",
    ])

    add_heading(doc, "Installation", level=2)
    add_numbered(doc, [
        "Open terminal in project root (C:\\WeddingPlannerApp).",
        "Run: npm install",
        "Run production check: npx next build",
    ])

    add_heading(doc, "Development Server (Environment Note)", level=2)
    add_paragraph(doc, "In this project environment, direct next binary invocation is the most reliable command for local runtime:")
    add_bullets(doc, [
        "node node_modules/next/dist/bin/next dev -p 3000",
        "Open http://localhost:3000 in browser",
    ])

    add_heading(doc, "Build and Release", level=2)
    add_bullets(doc, [
        "Build: npx next build",
        "Start production: npx next start",
        "Validate route generation and static optimization after each major UI/data change",
    ])


def add_testing_and_troubleshooting(doc: Document):
    add_heading(doc, "Testing Checklist", level=1)
    add_bullets(doc, [
        "Run full build and ensure all pages are generated without error.",
        "Validate favorites/compare/checklist persistence by browser refresh.",
        "Verify calculator totals for different event combinations and guest counts.",
        "Confirm responsive layouts on mobile/tablet/desktop breakpoints.",
        "Check external image loading from allowed domains.",
    ])

    add_heading(doc, "Troubleshooting Guide", level=1)
    table = doc.add_table(rows=1, cols=3)
    hdr = table.rows[0].cells
    hdr[0].text = "Issue"
    hdr[1].text = "Likely Cause"
    hdr[2].text = "Resolution"

    rows = [
        ("Dev server fails with npm run dev", "CLI wrapper inconsistency in environment", "Use node node_modules/next/dist/bin/next dev -p 3000"),
        ("Missing user state after reload", "localStorage disabled/cleared", "Enable storage and verify keys in browser devtools"),
        ("Incorrect comparison count", "Compare list max limit reached", "Remove one venue then add another"),
        ("Image not rendering", "Invalid URL or blocked domain", "Check URL and next.config.js image domain policy"),
    ]

    for r in rows:
        row = table.add_row().cells
        for i, col in enumerate(r):
            row[i].text = col


def add_security_and_limits(doc: Document):
    add_heading(doc, "Security, Privacy, and Current Limits", level=1)
    add_bullets(doc, [
        "No backend/database is used; this is a client-driven data model with static datasets.",
        "User planning data is stored in browser localStorage and remains device-local.",
        "No authentication/authorization is currently implemented.",
        "Contact/inquiry actions rely on tel/mail links, not integrated CRM APIs.",
        "Data updates are code-driven and require deployment to publish changes.",
    ])


def add_roadmap(doc: Document):
    add_heading(doc, "Recommended Enhancements", level=1)
    add_bullets(doc, [
        "Add backend API + database for real-time venue/vendor updates.",
        "Introduce authenticated user accounts and cloud sync.",
        "Implement PDF quote generation from calculator scenarios.",
        "Add analytics funnel for shortlist-to-contact conversion.",
        "Create admin CMS for non-technical content updates.",
    ])


def add_diagrams_section(doc: Document, diagrams):
    add_heading(doc, "System Diagrams", level=1)
    for title, image_path in diagrams:
        doc.add_paragraph(title)
        doc.add_picture(str(image_path), width=Inches(6.8))
        doc.add_paragraph("")


def generate_manual():
    DIAGRAM_DIR.mkdir(parents=True, exist_ok=True)

    architecture = DIAGRAM_DIR / "architecture_overview.png"
    workflows = DIAGRAM_DIR / "user_workflows.png"
    state_flow = DIAGRAM_DIR / "context_persistence_flow.png"
    budget_flow = DIAGRAM_DIR / "calculator_pipeline.png"

    diagram_architecture(architecture)
    diagram_user_workflows(workflows)
    diagram_state_flow(state_flow)
    diagram_budget_pipeline(budget_flow)

    doc = Document()

    write_title(doc)

    add_heading(doc, "Executive Summary", level=1)
    add_paragraph(doc, "Lahore Shaadi is a luxury wedding planning web application focused on Lahore venues and related planning workflows. The platform provides discovery, comparison, budgeting, checklist management, inspiration, and vendor coordination in a single Next.js application.")

    add_heading(doc, "Project Scope", level=1)
    add_bullets(doc, [
        "Venue discovery and filtering for premium marquees",
        "Detailed venue profile pages with packages and amenities",
        "Favorites and side-by-side comparison workflows",
        "Advanced multi-event budget calculator",
        "Planning checklist with persistent progress and countdown",
        "Vendor, gallery, and testimonial support workflows",
    ])

    add_heading(doc, "Technology Stack", level=1)
    add_tech_stack_table(doc)

    add_heading(doc, "Route Map", level=1)
    add_routes_table(doc)

    add_context_model(doc)
    add_workflows(doc)
    add_diagrams_section(doc, [
        ("Diagram 1 - Layered Architecture", architecture),
        ("Diagram 2 - Primary User Workflows", workflows),
        ("Diagram 3 - Context and Persistence Data Flow", state_flow),
        ("Diagram 4 - Budget Calculator Pipeline", budget_flow),
    ])
    add_devops(doc)
    add_testing_and_troubleshooting(doc)
    add_security_and_limits(doc)
    add_roadmap(doc)

    add_heading(doc, "Appendix: Key Project Modules", level=1)
    add_bullets(doc, [
        "pages/: route-level screens and workflow orchestration",
        "components/: reusable UI blocks (cards, modals, buttons, timers)",
        "context/WeddingContext.js: global state and persistence bridge",
        "data/: static domain models and utility selectors",
        "styles/globals.css + tailwind.config.js: design system and theme tokens",
    ])

    doc.save(OUT_DOC)

    print(f"Manual created: {OUT_DOC}")
    print(f"Diagrams created in: {DIAGRAM_DIR}")


if __name__ == "__main__":
    generate_manual()
